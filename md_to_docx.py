import argparse
import re
from pathlib import Path

import mistletoe
from mistletoe import HTMLRenderer
from bs4 import BeautifulSoup, Tag
from bs4.element import PageElement, NavigableString
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Length
from docx.table import Table
from docx.text.paragraph import Paragraph

class Pct(Length):
    """
    Represents a percentage-based width for Word tables.
    100% = 5000 units (1% = 50 units).
    """
    _EMUS_PER_PCT = 50  # Word uses 50 units per 1%

    def __new__(cls, percent):
        emu = int(percent * cls._EMUS_PER_PCT)
        return super().__new__(cls, emu)

    @property
    def pct(self):
        """Return the percentage value (0-100) from internal units."""
        return self.emu / self._EMUS_PER_PCT


class MarkdownToDocxConverter:
    def __init__(self, markdown_file: str | Path, docx_file: str | Path,
                 css_path: str | Path | None = None, embed_css_path: str | Path | None = None):
        self.markdown_file = Path(markdown_file)
        self.docx_file = Path(docx_file)
        self.css_path = Path(css_path) if css_path else None
        self.embed_css_path = Path(embed_css_path) if embed_css_path else None
        self.document = Document()

    def convert(self):
        self._validate_css_sources()

        with open(self.markdown_file, 'r', encoding='utf-8') as file:
            markdown_text = file.read()

        # Convert Markdown to HTML using mistletoe
        html_content = mistletoe.markdown(markdown_text, HTMLRenderer)

        soup = self._build_html_document(html_content)

        # Save the document as a HTML file for debugging purposes
        self._write_debug_html(soup)

        # Process the HTML elements and add them to the DOCX document
        parse_root = soup.body if soup.body else soup
        self._parse_html(parse_root)

        # Save the DOCX file
        self.document.save(str(self.docx_file))

    def _validate_css_sources(self):
        for css_source in (self.css_path, self.embed_css_path):
            if css_source and not css_source.is_file():
                raise FileNotFoundError(f"CSS file not found: {css_source}")

    def _build_html_document(self, html_fragment: str) -> BeautifulSoup:
        """Wrap raw HTML in a full document and inject optional CSS."""
        soup = BeautifulSoup('', 'html.parser')
        html_tag = soup.new_tag('html')
        head_tag = soup.new_tag('head')
        body_tag = soup.new_tag('body')

        fragment_soup = BeautifulSoup(html_fragment, 'html.parser')
        for child in fragment_soup.contents:
            body_tag.append(child)

        self._apply_css_to_head(soup, head_tag)

        html_tag.append(head_tag)
        html_tag.append(body_tag)
        soup.append(html_tag)
        return soup

    def _apply_css_to_head(self, soup: BeautifulSoup, head_tag: Tag):
        if self.css_path:
            link_tag = soup.new_tag('link', rel='stylesheet', href=self.css_path.as_posix())
            head_tag.append(link_tag)
        elif self.embed_css_path:
            style_tag = soup.new_tag('style')
            style_tag.string = self.embed_css_path.read_text(encoding='utf-8')
            head_tag.append(style_tag)

    def _write_debug_html(self, soup: BeautifulSoup):
        debug_html_path = self.docx_file.with_suffix('.html')
        debug_html_path.write_text(soup.prettify(), encoding='utf-8')

    def _parse_html(self, soup: BeautifulSoup):
        """Traverse all elements in the HTML DOM recursively"""
        for element in soup.children:
            if isinstance(element, Tag):
                self._process_element(element)

    def _process_element(self, element: PageElement, paragraph: Paragraph | None = None, level: int = 0):
        """Process HTML elements and convert them to DOCX content"""
        if isinstance(element, Tag):
            return self._process_tag_element(element, paragraph, level)
        elif isinstance(element, NavigableString):
            return self._process_text_element(element, paragraph)
        
        return paragraph

    def _process_tag_element(self, element: Tag, paragraph: Paragraph | None, level: int):
        """Process HTML tag elements"""
        # Block-level elements that create new paragraphs
        if self._is_block_element(element):
            return self._process_block_element(element, level)
        
        # Inline elements that require an existing paragraph
        elif self._is_inline_element(element):
            return self._process_inline_element(element, paragraph, level)
        
        # Unknown tags - process their children
        else:
            return self._process_unknown_element(element, paragraph, level)

    def _is_block_element(self, element: Tag) -> bool:
        """Check if element is a block-level element"""
        block_elements = ['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'h7', 'h8', 'h9', 
                         'p', 'ul', 'ol', 'table', 'blockquote', 'hr', 'pre', 'div', 'img']
        return element.name in block_elements

    def _is_inline_element(self, element: Tag) -> bool:
        """Check if element is an inline element"""
        inline_elements = ['strong', 'b', 'em', 'i', 'code', 'a', 'br']
        return element.name in inline_elements

    def _process_block_element(self, element: Tag, level: int):
        """Process block-level elements"""
        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'h7', 'h8', 'h9']:
            heading_level = int(element.name[1])
            return self._add_heading(element, level=heading_level)
        elif element.name == 'p':
            return self._add_paragraph(element)
        elif element.name == 'ul':
            self._add_list(element, ordered=False, level=level)
            return None
        elif element.name == 'ol':
            self._add_list(element, ordered=True, level=level)
            return None
        elif element.name == 'table':
            self._add_table(element)
            return None
        elif element.name == 'blockquote':
            return self._add_paragraph(element, style='Quote')
        elif element.name == 'hr':
            return self._add_horizontal_rule()
        elif element.name == 'pre':
            return self._add_preformatted_text(element)
        elif element.name == 'div':
            return self._process_div_element(element, level)
        elif element.name == 'img':
            return self._add_image_placeholder(element)

    def _process_inline_element(self, element: Tag, paragraph: Paragraph | None, level: int):
        """Process inline elements"""
        if paragraph is None:
            paragraph = self.document.add_paragraph(style='Normal')

        if element.name in ['strong', 'b']:
            self._process_formatted_inline_element(element, paragraph, level, bold=True)
        elif element.name in ['em', 'i']:
            self._process_formatted_inline_element(element, paragraph, level, italic=True)
        elif element.name == 'code':
            self._process_formatted_inline_element(element, paragraph, level, font_name='Courier New')
        elif element.name == 'a':
            self._process_link_element(element, paragraph, level)
        elif element.name == 'br':
            self._process_line_break(paragraph)

        return paragraph

    def _process_formatted_inline_element(self, element: Tag, paragraph: Paragraph, level: int, 
                                        bold: bool = False, italic: bool = False, font_name: str | None = None):
        """Process inline elements with formatting (bold, italic, code)"""
        for child in element.children:
            if isinstance(child, PageElement):
                if isinstance(child, NavigableString):
                    text = self._normalize_text_content(str(child))
                    if text:
                        run = paragraph.add_run(text)
                        if bold:
                            run.bold = True
                        if italic:
                            run.italic = True
                        if font_name:
                            run.font.name = font_name
                else:
                    self._process_element(child, paragraph, level)

    def _process_link_element(self, element: Tag, paragraph: Paragraph, level: int):
        """Process link elements"""
        for child in element.children:
            if isinstance(child, PageElement):
                self._process_element(child, paragraph, level)

    def _process_line_break(self, paragraph: Paragraph):
        """Process line break elements"""
        if paragraph and paragraph.style and 'List' in str(paragraph.style.name):
            paragraph.add_run(" ")  # Convert to space in list items
        else:
            paragraph.add_run("\n")

    def _process_unknown_element(self, element: Tag, paragraph: Paragraph | None, level: int):
        """Process unknown HTML tags by processing their children"""
        if paragraph is None:
            paragraph = self.document.add_paragraph(style='Normal')
        
        for child in element.children:
            if isinstance(child, PageElement):
                self._process_element(child, paragraph, level)
        
        return paragraph

    def _process_text_element(self, element: NavigableString, paragraph: Paragraph | None):
        """Process text content (NavigableString)"""
        if paragraph is None:
            paragraph = self.document.add_paragraph(style='Normal')

        text = self._normalize_text_content(str(element))
        if text:
            self._add_run(text, paragraph)

        return paragraph

    def _normalize_text_content(self, text: str) -> str:
        """Normalize whitespace in text content"""
        if text and text.strip():
            # Replace newlines, tabs, and multiple whitespace with single spaces
            return re.sub(r'\s+', ' ', text)
        return ""

    def _process_div_element(self, element: Tag, level: int):
        """Process div elements by processing their children"""
        for child in element.children:
            if isinstance(child, PageElement):
                self._process_element(child, None, level)
        return None

    def _add_horizontal_rule(self):
        """Add a horizontal rule"""
        p = self.document.add_paragraph(style='Normal')
        p.add_run('\n' + '-' * 50 + '\n')
        return p

    def _add_preformatted_text(self, element: Tag):
        """Add preformatted text"""
        p = self.document.add_paragraph(style='Quote')
        p.add_run(element.get_text())
        return p

    def _add_image_placeholder(self, element: Tag):
        """Add image placeholder text"""
        p = self.document.add_paragraph(style='Normal')
        p.add_run(f"[Image: {element.get('src', 'unknown')}]")
        return p

    def _add_run(self, text: str | None, paragraph: Paragraph, 
                font_name: str | None = None, bold: bool = False, italic: bool = False, 
                underline: bool = False):
        """Add a text run to a paragraph with optional formatting"""
        run = paragraph.add_run(text)
        if font_name:
            run.font.name = font_name
        run.bold = bold
        run.italic = italic
        run.underline = underline
        return run

    def _add_paragraph(self, element: Tag, style: str | None = None, level: int = 0):
        """Add a paragraph and process its children"""
        paragraph = self.document.add_paragraph(style=style)

        for child in element.children:
            if isinstance(child, PageElement):
                self._process_element(child, paragraph, level)

        return paragraph

    def _add_heading(self, element: Tag, level: int):
        """Add a heading with the specified level"""
        heading_style = f'Heading {min(level, 9)}'  # Word supports up to Heading 9
        paragraph = self.document.add_paragraph(style=heading_style)
        
        for child in element.children:
            if isinstance(child, PageElement):
                self._process_element(child, paragraph, level)
        
        return paragraph

    def _add_list(self, element: Tag, ordered: bool, level: int = 0):
        """Add a list (ordered or unordered) with proper nesting"""
        for child in element.children:
            if isinstance(child, Tag) and child.name == 'li':
                self._add_list_item(child, ordered, level)
            elif isinstance(child, NavigableString):
                continue  # Skip whitespace-only strings between list items

    def _add_list_item(self, li_element: Tag, ordered: bool, level: int):
        """Add a single list item"""
        # Determine list style based on type and level
        if ordered:
            list_style = 'List Number' if level == 0 else f'List Number {level + 1}' if level < 9 else 'List Number'
        else:
            list_style = 'List Bullet' if level == 0 else f'List Bullet {level + 1}' if level < 9 else 'List Bullet'
        
        paragraph = self.document.add_paragraph(style=list_style)
        
        # Process the contents of the list item
        for li_child in li_element.children:
            if isinstance(li_child, PageElement):
                if isinstance(li_child, NavigableString) and not str(li_child).strip():
                    continue  # Skip whitespace-only strings
                
                # Handle nested lists
                if isinstance(li_child, Tag) and li_child.name in ['ul', 'ol']:
                    nested_ordered = li_child.name == 'ol'
                    self._add_list(li_child, ordered=nested_ordered, level=level+1)
                else:
                    self._process_element(li_child, paragraph, level=level+1)

    def _add_table(self, element: Tag):
        """Add a table with proper formatting"""
        rows = element.find_all('tr')
        if not rows:
            return

        table = self.document.add_table(rows=len(rows), cols=len(rows[0].find_all(['td', 'th'])))
        table.style = 'Medium Shading 1 Accent 1'  # Set a default table style

        for row_index, row_element in enumerate(rows):
            for col_index, cell in enumerate(row_element.find_all(['td', 'th'])):
                self._populate_table_cell(table, row_index, col_index, cell)

        # For all columns except the last, set the width to Lengh = 0
        for col_index in range(len(table.columns) - 1):
            col = table.columns[col_index]
            for cell in col.cells:
                cell.width = Length(0)

        # Set table properties to 100% width and centered
        table.autofit = True  # Enable auto-fit to adjust column widths based on content
        self._set_table_width(table, Pct(100))  # Set table width to 100% of the page width

    def _populate_table_cell(self, table, row_index: int, col_index: int, cell_element):
        """Populate a single table cell with content and formatting"""
        cell_text = cell_element.get_text()
        table_cell = table.cell(row_index, col_index)
        table_cell.text = cell_text

    def _set_table_width(self, table: Table, width: Length | None):
        tbl = table._tbl
        tblPr = tbl.tblPr

        tblW = OxmlElement('w:tblW')
        
        # Set width type based on Length subclass
        if isinstance(width, Pct):
            tblW.set(qn('w:type'), 'pct')  # Use percentage
            tblW.set(qn('w:w'), str(int(width.pct * Pct._EMUS_PER_PCT)))  # Convert to Word's percentage units
        elif isinstance(width, Length):
            tblW.set(qn('w:type'), 'dxa')  # Use twips (1/20th of a point)
            tblW.set(qn('w:w'), str(width.twips))  # Set the width in twips
        elif not width:
            tblW.set(qn('w:type'), 'auto')  # Auto width
            tblW.set(qn('w:w'), '0')  # Set to 0 for auto width

        tblPr.append(tblW)

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Convert Markdown files to DOCX.")
    parser.add_argument("markdown_file", help="Input Markdown file")
    parser.add_argument("docx_file", nargs="?", help="Output DOCX file (defaults to <input>.docx)")

    css_group = parser.add_mutually_exclusive_group()
    css_group.add_argument("--css", dest="css_path", help="Path to a CSS file to link in the debug HTML output.")
    css_group.add_argument("--embed-css", dest="embed_css_path", help="Path to a CSS file to embed directly into the debug HTML output.")

    args = parser.parse_args()

    markdown_path = Path(args.markdown_file)
    if args.docx_file:
        docx_path = Path(args.docx_file)
    else:
        base = markdown_path.with_suffix('')
        docx_path = base.with_suffix('.docx')
        count = 1
        while docx_path.exists():
            docx_path = base.with_name(f"{base.name}_{count}").with_suffix('.docx')
            count += 1

    converter = MarkdownToDocxConverter(markdown_path, docx_path, css_path=args.css_path, embed_css_path=args.embed_css_path)
    converter.convert()